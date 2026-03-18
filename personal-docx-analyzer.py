#!/usr/bin/env python3

"""
🤖 Zalo Personal Messages Analyzer - DOCX Report
- Quét tin nhắn cá nhân Zalo
- Tạo báo cáo Word (.docx)
- Tối ưu hóa cho xem trên điện thoại
- Lưu tại: ~/zalo-ai-bot/reports/
"""

import subprocess
import json
import sys
import os
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# ZALO MCP WRAPPER
# ============================================================

class ZaloMCP:
    """Gọi MCP tools để interact với Zalo"""
    
    @staticmethod
    def get_personal_threads() -> List[Dict[str, Any]]:
        """Lấy danh sách cuộc trò chuyện CÁ NHÂN (type=0)"""
        try:
            cmd = "zalo-agent conv recent --json"
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=10)
            
            if not result.stdout:
                return []
            
            all_threads = json.loads(result.stdout)
            
            # Lọc chỉ cuộc trò chuyện cá nhân (typeFlag=0 hoặc type='User')
            personal = [t for t in all_threads if t.get('typeFlag') == 0 or t.get('type') == 'User']
            
            return personal
        
        except Exception as e:
            print(f"❌ Lỗi lấy cuộc trò chuyện: {e}")
            return []

# ============================================================
# PERSONAL DATA COLLECTOR
# ============================================================

class PersonalDataCollector:
    """Quét dữ liệu tin nhắn cá nhân"""
    
    @staticmethod
    def collect_personal_messages(limit: int = 20) -> Dict[str, Any]:
        """Quét tin nhắn từ cuộc trò chuyện cá nhân"""
        
        print("📊 Đang quét tin nhắn cá nhân từ Zalo...")
        print("─" * 60)
        
        # Lấy danh sách cuộc trò chuyện cá nhân
        personal_threads = ZaloMCP.get_personal_threads()
        
        if not personal_threads:
            print("❌ Không tìm thấy cuộc trò chuyện cá nhân nào")
            return {"success": False, "error": "No personal threads found"}
        
        print(f"✓ Tìm thấy {len(personal_threads)} cuộc trò chuyện cá nhân")
        
        # Quét từng cuộc trò chuyện
        collected_data = {}
        
        for i, thread in enumerate(personal_threads[:limit], 1):
            thread_id = thread.get('threadId')
            thread_name = thread.get('name', 'Unknown')
            
            print(f"\n[{i}/{min(len(personal_threads), limit)}] {thread_name}")
            
            collected_data[thread_id] = {
                "name": thread_name,
                "threadId": thread_id,
                "type": "personal",
                "typeFlag": thread.get('typeFlag', 0),
                "lastActive": thread.get('lastActive', ''),
                "memberCount": thread.get('memberCount', 1),
            }
            
            print(f"   ✓ Hoạt động gần nhất: {thread.get('lastActive', 'N/A')}")
        
        print("\n" + "─" * 60)
        print(f"✅ Quét xong! Tổng {len(collected_data)} cuộc trò chuyện cá nhân")
        
        return {
            "success": True,
            "timestamp": datetime.now().isoformat(),
            "conversationCount": len(collected_data),
            "data": collected_data
        }

# ============================================================
# ANALYSIS ENGINE
# ============================================================

class AnalysisEngine:
    """Phân tích dữ liệu tin nhắn"""
    
    @staticmethod
    def analyze_conversations(data: Dict[str, Any]) -> Dict[str, Any]:
        """Phân tích và tạo insights"""
        
        conversations = data.get('data', {})
        
        if not conversations:
            return {"error": "Không có dữ liệu để phân tích."}
        
        # Phân loại theo hoạt động
        sorted_convs = sorted(
            conversations.values(),
            key=lambda x: x.get('lastActive', ''),
            reverse=True
        )
        
        return {
            "total_conversations": len(conversations),
            "sorted_conversations": sorted_convs,
            "insights": {
                "recent_active": sorted_convs[0] if sorted_convs else None,
                "oldest_contact": sorted_convs[-1] if sorted_convs else None,
            }
        }

# ============================================================
# DOCX REPORT GENERATOR
# ============================================================

class DocxReportGenerator:
    """Tạo báo cáo Word (.docx) tối ưu cho mobile"""
    
    @staticmethod
    def add_heading_style(doc, text: str, level: int = 1):
        """Thêm heading"""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    @staticmethod
    def add_table_with_data(doc, headers: List[str], rows: List[List[str]]):
        """Thêm bảng dữ liệu"""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Tô màu header
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # Rows
        for row_idx, row_data in enumerate(rows, 1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
    
    @staticmethod
    def generate(data: Dict[str, Any], analysis: Dict[str, Any]) -> Document:
        """Tạo document Word"""
        
        doc = Document()
        
        # ============================================================
        # TITLE PAGE
        # ============================================================
        
        title = doc.add_heading('📊 Báo cáo Phân tích\nTin Nhắn Zalo Cá Nhân', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0].font
        title_format.size = Pt(28)
        title_format.bold = True
        title_format.color.rgb = RGBColor(0, 102, 204)
        
        # Thông tin cơ bản
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp = data.get('timestamp', datetime.now().isoformat())
        conv_count = data.get('conversationCount', 0)
        
        info_text = f"Ngày: {datetime.fromisoformat(timestamp).strftime('%d/%m/%Y %H:%M:%S')}\n"
        info_text += f"Tổng cuộc trò chuyện: {conv_count}\n"
        info_text += f"Múi giờ: Asia/Ho_Chi_Minh"
        
        info_para.text = info_text
        info_para_format = info_para.runs[0].font
        info_para_format.size = Pt(12)
        
        # Divider
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
        
        # ============================================================
        # SUMMARY SECTION
        # ============================================================
        
        doc.add_heading('📈 Tóm tắt', 1)
        
        summary_data = [
            ['Cuộc trò chuyện cá nhân', str(conv_count)],
            ['Loại tin nhắn', 'Cuộc trò chuyện riêng'],
            ['Timezone', 'Asia/Ho_Chi_Minh'],
        ]
        
        DocxReportGenerator.add_table_with_data(
            doc, 
            ['Thông tin', 'Giá trị'],
            summary_data
        )
        
        doc.add_paragraph()
        
        # ============================================================
        # ANALYSIS SECTION
        # ============================================================
        
        doc.add_heading('🔍 Phân tích Chi tiết', 1)
        
        conversations = data.get('data', {})
        sorted_convs = analysis.get('sorted_conversations', [])
        
        if sorted_convs:
            doc.add_paragraph(
                f"Được sắp xếp theo hoạt động gần nhất:\n",
                style='List Bullet'
            )
            
            # Tạo bảng danh sách
            conv_rows = []
            for i, conv in enumerate(sorted_convs, 1):
                conv_rows.append([
                    str(i),
                    conv.get('name', 'Unknown'),
                    conv.get('lastActive', 'N/A'),
                ])
            
            DocxReportGenerator.add_table_with_data(
                doc,
                ['#', 'Tên liên hệ', 'Hoạt động gần nhất'],
                conv_rows
            )
        
        doc.add_paragraph()
        
        # ============================================================
        # INSIGHTS SECTION
        # ============================================================
        
        doc.add_heading('💡 Insights', 1)
        
        insights = analysis.get('insights', {})
        recent = insights.get('recent_active', {})
        oldest = insights.get('oldest_contact', {})
        
        if recent:
            doc.add_paragraph(
                f"🟢 Cuộc trò chuyện gần nhất: {recent.get('name', 'N/A')}",
                style='List Bullet'
            )
        
        if oldest:
            doc.add_paragraph(
                f"🔵 Liên hệ cũ nhất: {oldest.get('name', 'N/A')}",
                style='List Bullet'
            )
        
        doc.add_paragraph(
            f"📱 Tổng cộng: {conv_count} cuộc trò chuyện cá nhân",
            style='List Bullet'
        )
        
        doc.add_paragraph()
        
        # ============================================================
        # DETAILED LIST SECTION
        # ============================================================
        
        doc.add_heading('📋 Danh sách Chi tiết', 1)
        
        for i, conv in enumerate(sorted_convs, 1):
            # Sub-heading
            sub_heading = doc.add_heading(
                f"{i}. {conv.get('name', 'Unknown')}",
                level=2
            )
            
            # Chi tiết
            details = [
                ['Thread ID', conv.get('threadId', 'N/A')],
                ['Hoạt động cuối', conv.get('lastActive', 'N/A')],
                ['Loại', 'Cuộc trò chuyện riêng'],
            ]
            
            DocxReportGenerator.add_table_with_data(
                doc,
                ['Thuộc tính', 'Giá trị'],
                details
            )
            
            doc.add_paragraph()  # Space
        
        # ============================================================
        # RECOMMENDATIONS SECTION
        # ============================================================
        
        doc.add_heading('✅ Gợi ý hành động', 1)
        
        doc.add_paragraph(
            'Kiểm tra cuộc trò chuyện gần nhất',
            style='List Number'
        )
        doc.add_paragraph(
            'Liên hệ với những người chưa nói chuyện lâu',
            style='List Number'
        )
        doc.add_paragraph(
            'Sắp xếp lại danh sách bạn bè theo ưu tiên',
            style='List Number'
        )
        doc.add_paragraph(
            'Dọn dẹp những cuộc trò chuyện cũ nếu cần',
            style='List Number'
        )
        
        doc.add_paragraph()
        
        # ============================================================
        # FOOTER SECTION
        # ============================================================
        
        doc.add_paragraph("_" * 60)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = footer.add_run(
            f"📱 Báo cáo Phân tích Zalo Cá Nhân\n"
            f"Tạo lúc: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n"
            f"Powered by Zalo AI Bot\n"
            f"Version 1.0"
        )
        footer_text.font.size = Pt(10)
        footer_text.font.color.rgb = RGBColor(128, 128, 128)
        
        return doc
    
    @staticmethod
    def save(doc: Document, filename: str = None) -> str:
        """Lưu document"""
        
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"personal_report_{timestamp}.docx"
        
        # Tạo thư mục reports
        os.makedirs("reports", exist_ok=True)
        
        filepath = os.path.join("reports", filename)
        doc.save(filepath)
        
        return filepath

# ============================================================
# MAIN PIPELINE
# ============================================================

def main():
    """Main entry point - Complete pipeline"""
    
    print("\n" + "═" * 60)
    print("🤖 ZALO PERSONAL ANALYZER - DOCX REPORT".center(60))
    print("═" * 60 + "\n")
    
    # Step 1: Quét dữ liệu
    print("📊 STEP 1: Quét dữ liệu tin nhắn cá nhân")
    print("─" * 60)
    
    data = PersonalDataCollector.collect_personal_messages(limit=20)
    
    if not data.get('success'):
        print(f"❌ Lỗi: {data.get('error')}")
        sys.exit(1)
    
    # Step 2: Phân tích
    print("\n🧠 STEP 2: Phân tích dữ liệu")
    print("─" * 60)
    
    analysis = AnalysisEngine.analyze_conversations(data)
    print(f"✓ Phân tích {data.get('conversationCount', 0)} cuộc trò chuyện")
    
    # Step 3: Tạo báo cáo Word
    print("\n📝 STEP 3: Tạo báo cáo Word (.docx)")
    print("─" * 60)
    
    doc = DocxReportGenerator.generate(data, analysis)
    filepath = DocxReportGenerator.save(doc)
    
    print(f"✓ Báo cáo lưu tại: {filepath}")
    
    # Step 4: Hiển thị thông tin
    print("\n" + "═" * 60)
    print("✅ HOÀN THÀNH!".center(60))
    print("═" * 60)
    
    print(f"\n📁 Báo cáo: {filepath}")
    print(f"📊 Cuộc trò chuyện: {data.get('conversationCount', 0)}")
    print(f"📱 Định dạng: Word (.docx) - Tiện xem trên điện thoại")
    print(f"\n💡 Mở báo cáo:")
    print(f"   open {filepath}")
    print(f"\n📂 Thư mục báo cáo:")
    print(f"   ~/zalo-ai-bot/reports/")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⏹️ Bị dừng bởi người dùng")
        sys.exit(0)
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
